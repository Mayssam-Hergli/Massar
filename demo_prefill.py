"""Demo: parse a pitch deck → pre-fill & verify the founder's questionnaire answers.

Builds a real .docx pitch deck, runs it through the File Parser, maps its content to
the venture questionnaire keys, then cross-checks a founder's submitted answers against
the document (flagging matches, mismatches, and unanswered-but-suggested fields).

Run:  python demo_prefill.py
"""

from __future__ import annotations

import asyncio
import sys
from io import BytesIO

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from services.ms1_diagnostic.prefill import DocumentPrefillService

PITCH_DECK_LINES = [
    "GreenLoop — Pitch Deck",
    "Problem: industrial sites waste water and energy at scale.",
    "Solution: a fully automated, AI-driven water recycling platform (circular economy).",
    "Our proprietary technology has a patent pending and runs on solar power.",
    "Traction: we interviewed 24 customers and signed 2 letters of intent.",
    "We already have paying customers generating recurring revenue (ARR).",
    "Business model: SaaS subscription, pricing model defined at $499 / month.",
    "Product is live and in production, validated for product-market fit.",
    "Market opportunity: a $12 billion TAM with global expansion planned.",
    "We are raising a Series A round to scale internationally.",
]


def _build_pitch_deck_docx() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("GreenLoop", level=1)
    for line in PITCH_DECK_LINES:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# What the founder typed into the form — note two deliberate disagreements with the deck.
FOUNDER_ANSWERS = {
    "market_size": "medium",          # deck implies very_large ($12B TAM) → MISMATCH
    "has_paying_customers": False,    # deck says paying customers / ARR → MISMATCH
    "product_maturity": "product",    # agrees with deck
    "energy_source": "solar_wind",    # agrees with deck
    # ...the rest left blank → eligible for pre-fill
}


async def main() -> None:
    print("=" * 74)
    print("MS1 PRE-FILL & VERIFY — pitch deck → questionnaire")
    print("=" * 74)

    deck = _build_pitch_deck_docx()
    service = DocumentPrefillService()

    suggestions = await service.suggest_from_document(
        data=deck, filename="greenloop_deck.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    print(f"\nParsed the deck and derived {len(suggestions)} suggested answer(s):\n")
    print(f"  {'key':26}{'suggested':16}{'conf':>5}  evidence")
    for s in suggestions:
        print(f"  {s.key:26}{str(s.suggested_value):16}{s.confidence:>5}  \"{s.evidence[:46]}\"")

    print("\n--- Verifying the founder's submitted answers against the deck ---\n")
    verifications = service.verify(answers=FOUNDER_ANSWERS, suggestions=suggestions)
    icon = {"match": "OK  ", "mismatch": "FLAG", "prefill": "FILL"}
    for v in sorted(verifications, key=lambda x: x.status):
        line = (f"  [{icon[v.status]}] {v.key:26} answered={str(v.answered_value):8} "
                f"deck={str(v.suggested_value):14} ({v.status})")
        print(line)

    mismatches = [v for v in verifications if v.status == "mismatch"]
    prefills = [v for v in verifications if v.status == "prefill"]
    print(f"\nSummary: {sum(1 for v in verifications if v.status=='match')} match, "
          f"{len(mismatches)} to review, {len(prefills)} auto-fillable.")

    merged = service.prefilled_answers(answers=FOUNDER_ANSWERS, suggestions=suggestions)
    print(f"After pre-fill, the founder's form went from {len(FOUNDER_ANSWERS)} "
          f"to {len(merged)} answered keys (they can still override).")
    print("\nNote: mismatches are surfaced for the founder to confirm — MS1 still writes "
          "whatever they finally submit, so the MS2 contract stays clean.")


if __name__ == "__main__":
    asyncio.run(main())
